"""
CV Parser Service

Extracts text content from uploaded CV/Resume files.
Supports PDF and DOCX formats.

The extracted text is used in the system prompt so Gemini
can ask personalized interview questions based on the candidate's
experience, projects, and skills.
"""

import io
import re
import logging
from typing import Optional

logger = logging.getLogger(__name__)

# Optional imports for file parsing
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("pdfplumber not installed. PDF parsing disabled. Run: pip install pdfplumber")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. DOCX parsing disabled. Run: pip install python-docx")


class CVParser:
    """
    Parses CV/Resume files and extracts text content.
    
    Example usage:
        parser = CVParser()
        cv_text = parser.parse(file_bytes, "resume.pdf")
        # cv_text contains cleaned text from the CV
    """
    
    MAX_CHARS = 2000  # Limit to keep system prompt reasonable
    
    def parse(self, file_bytes: bytes, filename: str) -> str:
        """
        Parse a CV file and extract text content.
        
        Args:
            file_bytes: Raw bytes of the uploaded file
            filename: Original filename (used to detect format)
        
        Returns:
            Cleaned text content, truncated to MAX_CHARS
        """
        filename_lower = filename.lower()
        
        try:
            if filename_lower.endswith('.pdf'):
                return self._parse_pdf(file_bytes)
            elif filename_lower.endswith('.docx'):
                return self._parse_docx(file_bytes)
            elif filename_lower.endswith('.doc'):
                logger.warning("Legacy .doc format not supported. Please use .docx")
                return ""
            elif filename_lower.endswith('.txt'):
                return self._parse_txt(file_bytes)
            else:
                logger.warning(f"Unsupported file format: {filename}")
                return ""
                
        except Exception as e:
            logger.error(f"Failed to parse CV file {filename}: {e}")
            return ""
    
    def _parse_pdf(self, file_bytes: bytes) -> str:
        """Extract text from PDF file."""
        if not PDF_AVAILABLE:
            raise ImportError("pdfplumber not installed")
        
        text_parts = []
        
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        raw_text = "\n".join(text_parts)
        return self._clean_text(raw_text)
    
    def _parse_docx(self, file_bytes: bytes) -> str:
        """Extract text from DOCX file."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed")
        
        doc = Document(io.BytesIO(file_bytes))
        
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract from tables if present
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)
        
        raw_text = "\n".join(text_parts)
        return self._clean_text(raw_text)
    
    def _parse_txt(self, file_bytes: bytes) -> str:
        """Extract text from plain text file."""
        try:
            text = file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            text = file_bytes.decode('latin-1')
        
        return self._clean_text(text)
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.
        
        - Removes excessive whitespace
        - Removes control characters
        - Truncates to MAX_CHARS
        """
        # Remove control characters except newlines and tabs
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        # Normalize whitespace
        text = re.sub(r' +', ' ', text)  # Multiple spaces to single
        text = re.sub(r'\n{3,}', '\n\n', text)  # Multiple newlines to double
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        # Truncate if too long
        if len(text) > self.MAX_CHARS:
            # Try to truncate at a sentence boundary
            truncated = text[:self.MAX_CHARS]
            last_period = truncated.rfind('.')
            if last_period > self.MAX_CHARS * 0.8:
                truncated = truncated[:last_period + 1]
            text = truncated + "\n[CV truncated for brevity]"
        
        return text


# Singleton instance
_parser: Optional[CVParser] = None


def get_cv_parser() -> CVParser:
    """Get singleton CV parser instance."""
    global _parser
    if _parser is None:
        _parser = CVParser()
    return _parser


def parse_cv(file_bytes: bytes, filename: str) -> str:
    """Convenience function to parse a CV file."""
    return get_cv_parser().parse(file_bytes, filename)
